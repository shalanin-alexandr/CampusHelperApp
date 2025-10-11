import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

nltk.download('punkt', quiet=True)

# 📄 Чтение TXT
def read_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

# 📄 Чтение DOCX
def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    return '\n'.join(para.text for para in doc.paragraphs)

# 📤 Сохранение в DOCX
def save_docx(summary: list[str], filename: str) -> str:
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = os.path.join("generated", filename)

    doc = Document()
    heading = doc.add_heading("Сокращённый текст", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in summary:
        doc.add_paragraph(line)
    os.makedirs("generated", exist_ok=True)
    doc.save(filepath)
    return filepath

# 🧠 Сокращение текста
def summarize_text(text: str) -> list[str]:
    if not text.strip():
        return []

    sentences = nltk.sent_tokenize(text)
    if len(sentences) < 3:
        return sentences

    tf_idf = TfidfVectorizer().fit_transform(sentences)
    scores = tf_idf.sum(axis=1).A1

    top_n = max(4, len(sentences) // 4)
    top_idx = scores.argsort()[-top_n:][::-1]
    summary = [sentences[i] for i in sorted(top_idx)]
    return summary
