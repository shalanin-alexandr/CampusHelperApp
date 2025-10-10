import requests
import io
import re
import os
from docx import Document
from datetime import datetime, timedelta
import json

# Абсолютный путь для файла кэша
CACHE_FILE = os.path.join(os.path.dirname(__file__), "last_docx_url.txt")

weekday_map = {
    0: "Понедельник",
    1: "Вторник",
    2: "Среда",
    3: "Четверг",
    4: "Пятница",
    5: "Суббота",
    6: "Воскресенье"
}

def normalize_day(day):
    return day.strip().lower().replace("ё", "е") if isinstance(day, str) else ""

def normalize_group(g):
    return re.sub(r"\W+", "", g.strip().lower()) if isinstance(g, str) else ""

def fetch_latest_docx_url(page_url):
    try:
        response = requests.get(page_url)
        response.raise_for_status()
        html = response.text

        match = re.search(r'href="([^"]*Zamena_SAYT\.docx[^"]*)"', html)
        if match:
            return "http://www.bobruisk.belstu.by" + match.group(1)
        else:
            print("❌ Не удалось найти ссылку на DOCX.")
            return None
    except Exception as e:
        print(f"❌ Ошибка при получении ссылки: {e}")
        return None

def get_week_type_from_docx(doc):
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.lower())

    text_content = " ".join(full_text)

    if "верхняя неделя" in text_content:
        return "upper"
    elif "нижняя неделя" in text_content:
        return "lower"
    return None

def has_docx_url_changed(new_url, cache_file=CACHE_FILE):
    try:
        with open(cache_file, "r") as f:
            old_url = f.read().strip()
    except FileNotFoundError:
        old_url = ""

    if new_url != old_url:
        with open(cache_file, "w") as f:
            f.write(new_url)
        return True
    return False

def load_docx_from_url(url):
    response = requests.get(url)
    response.raise_for_status()
    return Document(io.BytesIO(response.content))

def get_day_from_table(table):
    for row in table.rows[:2]:
        for cell in row.cells:
            text = cell.text.strip().upper()
            match = re.search(r'(ПОНЕДЕЛЬНИК|ВТОРНИК|СРЕДА|ЧЕТВЕРГ|ПЯТНИЦА|СУББОТА|ВОСКРЕСЕНЬЕ)', text)
            if match:
                return match.group(1).capitalize()
    return None


def parse_schedule_table(table, target_group, day_label):
    schedule_data = {
        "group": target_group,
        "schedule": []
    }

    current_group = None
    group_found = False

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]

        if not any(cells):
            continue

        if len(cells) == 2 and cells[0] and cells[1]:
            group = cells[0]
            comment = cells[1]
            if normalize_group(group) == normalize_group(target_group):
                group_found = True
                schedule_data["schedule"].append({
                    "day": day_label,
                    "comment": comment
                })
            continue

        if cells[0] and not cells[0].startswith("-"):
            current_group = cells[0]
            group_found = normalize_group(current_group) == normalize_group(target_group)

        if not group_found:
            continue

        try:
            pair_number = cells[1] if len(cells) > 1 else None
            room = cells[2] if len(cells) > 2 else None
            subject_to = cells[3] if len(cells) > 3 else None
            teacher_to = cells[4] if len(cells) > 4 else None
            subject_from = cells[5] if len(cells) > 5 else None
            teacher_from = cells[7] if len(cells) > 6 else None

            schedule_data["schedule"].append({
                "day": day_label,
                "pair": str(pair_number).strip() if pair_number else None,
                "room": room,
                "from": {
                    "subject": subject_from,
                    "teacher": teacher_from
                },
                "to": {
                    "subject": subject_to,
                    "teacher": teacher_to
                }
            })
        except Exception as e:
            print(f"⚠️ Ошибка при обработке строки: {e}")
            continue

    return schedule_data

def get_docx_schedule(group_name, page_url="http://www.bobruisk.belstu.by/dnevnoe-otdelenie/raspisanie-zanyatiy-i-zvonkov-zamenyi#gsc.tab=0", doc_updated=False):
    docx_url = fetch_latest_docx_url(page_url)
    if not docx_url:
        return None

    try:
        doc = load_docx_from_url(docx_url)

        full_schedule = {
            "group": group_name,
            "schedule": []
        }

        week_type = get_week_type_from_docx(doc)
        full_schedule["week_type"] = week_type
        print(f"📌 Тип недели: {week_type}")

        # Собираем список дней из абзацев
        day_labels = []
        for para in doc.paragraphs:
            text = para.text.strip().upper()
            match = re.search(r'(ПОНЕДЕЛЬНИК|ВТОРНИК|СРЕДА|ЧЕТВЕРГ|ПЯТНИЦА|СУББОТА|ВОСКРЕСЕНЬЕ)', text)
            if match:
                day_labels.append(match.group(1).capitalize())

        print(f"📅 Найденные дни перед таблицами: {day_labels}")

        # Привязка дней к таблицам
        for i, table in enumerate(doc.tables):
            if i < len(day_labels):
                day_label = day_labels[i]
            else:
                # Если таблиц больше, чем дней — используем резерв
                target_day = datetime.now() + timedelta(days=1)
                if target_day.weekday() == 6: 
                    target_day += timedelta(days=1)
                day_label = weekday_map[target_day.weekday()]
                print(f"⚠️ День не найден для таблицы {i}, используем резерв: {day_label}")

            result = parse_schedule_table(table, group_name, day_label)
            full_schedule["schedule"].extend(result["schedule"])

        if full_schedule["schedule"]:
            return full_schedule
        else:
            print(f"⚠️ Группа {group_name} не найдена в документе.")
            return full_schedule

    except requests.exceptions.RequestException as e:
        print(f"❌ Ошибка при скачивании файла: {e}")
        return None
    except Exception as e:
        print(f"❌ Ошибка при обработке DOCX: {e}")
        return None


def get_available_replacement_days(doc_schedule):
    days_with_replacements = set()

    for item in doc_schedule.get("schedule", []):
        if "to" in item and item.get("day"):
            days_with_replacements.add(item["day"].strip())

    return sorted(days_with_replacements)

if __name__ == '__main__':
    MY_GROUP = "РС02-24"
    DOC_PAGE_URL = "http://www.bobruisk.belstu.by/dnevnoe-otdelenie/raspisanie-zanyatiy-i-zvonkov-zamenyi"
    
    DOCX_URL = fetch_latest_docx_url(DOC_PAGE_URL)
    doc_updated = has_docx_url_changed(DOCX_URL) 
    
    schedule = get_docx_schedule(MY_GROUP, DOC_PAGE_URL, doc_updated) 

    if schedule:
        print("✅ Получено расписание замен:")
        print(json.dumps(schedule, indent=4, ensure_ascii=False))

        available_days = get_available_replacement_days(schedule)
        print("\n📌 Дни, для которых найдены замены:")
        print(available_days)
